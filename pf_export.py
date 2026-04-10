# ─────────────────────────────────────────────────────────────────────────────
# pf_export.py — Export Project Finance Contract Minutes to DOCX / Markdown
# ─────────────────────────────────────────────────────────────────────────────
import io
import re
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from pf_models import GeneratedMinutes, ContractDocument


# ─── Helpers ─────────────────────────────────────────────────────────────────

PLACEHOLDER_RE = re.compile(r"(\[●\]|\[[A-Z][^\]]*\])")


def _add_runs_with_placeholders(paragraph, text: str, base_size_pt: int = 11):
    """Split text on placeholders and render placeholders in bold red."""
    parts = PLACEHOLDER_RE.split(text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.size = Pt(base_size_pt)
        run.font.name = "Calibri"
        if PLACEHOLDER_RE.fullmatch(part):
            run.bold = True
            run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)  # red


def _add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)


def _add_subtitle(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


def _add_section_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def _add_clause_heading(doc: Document, number: int, title: str):
    p = doc.add_paragraph()
    run = p.add_run(f"{number}. {title}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def _add_body_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    _add_runs_with_placeholders(p, text)


def _add_party(doc: Document, idx: int, name: str, role: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    run_idx = p.add_run(f"({idx}) ")
    run_idx.font.size = Pt(11)
    run_idx.font.name = "Calibri"
    run_name = p.add_run(name)
    run_name.bold = True
    run_name.font.size = Pt(11)
    run_name.font.name = "Calibri"
    run_role = p.add_run(f" (the '{role}')")
    run_role.font.size = Pt(11)
    run_role.font.name = "Calibri"


# ─── Document writers ────────────────────────────────────────────────────────

def _write_contract_to_doc(doc: Document, contract: ContractDocument,
                           project_name: str, closing_date: str = ""):
    """Append a single contract document to a Document object."""
    _add_title(doc, contract.type.upper())
    if closing_date:
        _add_subtitle(doc, f"Dated {closing_date}")
    if project_name:
        _add_subtitle(doc, project_name)

    doc.add_paragraph()  # blank line

    # Parties
    _add_section_heading(doc, "PARTIES")
    if contract.parties:
        for i, p in enumerate(contract.parties, start=1):
            _add_party(doc, i, p["name"], p["role"])
    else:
        _add_body_paragraph(doc, "[No parties identified]")

    # Recitals (boilerplate)
    _add_section_heading(doc, "RECITALS")
    _add_body_paragraph(doc,
        f"(A) The Borrower is undertaking the Project described herein.")
    _add_body_paragraph(doc,
        f"(B) The Borrower has requested the Lenders to provide financing in respect of the Project on the terms set out in this Agreement.")
    _add_body_paragraph(doc,
        f"(C) The Lenders have agreed to provide such financing on the terms and subject to the conditions set out herein.")

    p = doc.add_paragraph()
    run = p.add_run("NOW THIS AGREEMENT WITNESSES as follows:")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # Clauses
    for cl in contract.clauses:
        _add_clause_heading(doc, cl.number, cl.title)
        _add_body_paragraph(doc, cl.text)

    # Signature block
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("EXECUTED as a deed by the parties hereto on the date first written above.")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    doc.add_paragraph()
    if contract.parties:
        for party in contract.parties:
            sig = doc.add_paragraph()
            sig.add_run("_______________________________").font.name = "Calibri"
            sig2 = doc.add_paragraph()
            run2 = sig2.add_run(f"For and on behalf of {party['name']} ({party['role']})")
            run2.font.size = Pt(11)
            run2.font.name = "Calibri"
            doc.add_paragraph()


# ─── Public API ──────────────────────────────────────────────────────────────

def export_minutes_to_docx(minutes: GeneratedMinutes,
                            doc_type: Optional[str] = None,
                            closing_date: str = "") -> bytes:
    """Export GeneratedMinutes to a .docx file as bytes.

    If doc_type is None, exports ALL documents in one file with page breaks.
    If doc_type is specified, exports only that single document.
    """
    doc = Document()

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)

    docs_to_export = (
        [d for d in minutes.documents if d.type == doc_type] if doc_type
        else minutes.documents
    )

    if not docs_to_export:
        # Empty fallback
        _add_title(doc, "NO DOCUMENTS GENERATED")
        _add_body_paragraph(doc, "No contract documents matched the selected criteria.")
    else:
        for i, contract in enumerate(docs_to_export):
            if i > 0:
                doc.add_page_break()
            _write_contract_to_doc(doc, contract, minutes.project_name, closing_date)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def export_minutes_to_markdown(minutes: GeneratedMinutes,
                                doc_type: Optional[str] = None) -> str:
    """Combine all (or one) contract documents into a single markdown string."""
    docs_to_export = (
        [d for d in minutes.documents if d.type == doc_type] if doc_type
        else minutes.documents
    )
    if not docs_to_export:
        return "# No documents generated"

    parts = []
    for d in docs_to_export:
        parts.append(d.markdown)
        parts.append("\n\n---\n\n")
    return "".join(parts)
